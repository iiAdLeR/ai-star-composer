# -*- coding: utf-8 -*-
"""Tek seferlik PAT (Planlama-Analiz-Tasarım) Word belgesi üretir."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _set_cell_text(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)


def main() -> None:
    out = Path(__file__).resolve().parent / "PAT_Mezuniyet_AI_Star_Composer.docx"
    d = Document()

    # Kapak benzeri başlık
    t = d.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MEZUNİYET PROJESİ\nPAT DOKÜMANI")
    r.bold = True
    r.font.size = Pt(18)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Planlama · Analiz · Tasarım\n\n"
        "Proje Adı: AI Star Composer\n"
        "Konu: Gezegen Keşif Verilerine Dayalı Yapay Zekâ Destekli Müzik Üretimi\n"
    ).font.size = Pt(12)

    d.add_paragraph()
    meta = d.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Hazırlanma Tarihi: Nisan 2026\n"
        "Ders: Mezuniyet Projesi\n\n"
        "Öğrenci(ler): [Ad Soyad — doldurunuz]\n"
        "Öğrenci No: [—]\n"
        "Danışman: [Unvan Ad Soyad — doldurunuz]\n"
        "Bölüm / Kurum: [—]\n"
    ).font.size = Pt(11)

    d.add_page_break()

    d.add_heading("İçindekiler", level=1)
    for line in (
        "1. Giriş ve Proje Özeti",
        "2. Planlama (Planlama)",
        "3. Analiz",
        "4. Tasarım",
        "5. PAT Kapsamında Teslim Özeti",
        "Kaynakça",
    ):
        d.add_paragraph(line, style="List Bullet")

    d.add_page_break()

    # --- 1. Giriş ---
    d.add_heading("1. Giriş ve Proje Özeti", level=1)
    d.add_paragraph(
        "Bu belge, mezuniyet projesi kapsamında vize öncesi tamamlanması istenen "
        "Planlama (Planlama), Analiz ve Tasarım (PAT) aşamalarını tek çatı altında "
        "özetlemektedir. Proje; NASA ve benzeri kaynaklardan elde edilen gezegen "
        "hareket/keşif verilerini müzikal parametrelere dönüştürerek, kullanıcıya "
        "etkileşimli bir stüdyo ortamında müzik üretimi sunmayı hedeflemektedir."
    )
    d.add_paragraph(
        "Sistem; web tabanlı arayüz, arka uç API, sembolik müzik motoru, isteğe bağlı "
        "LSTM tabanlı öğrenme ile melodi zenginleştirmesi ve canlı akış senaryolarını "
        "kapsar. Bu dokümanda yazılımın tam kodu yerine gereksinimler, mimari kararlar "
        "ve tasarım çıktıları vurgulanmıştır."
    )

    # --- 2. Planlama ---
    d.add_heading("2. Planlama (Planlama)", level=1)

    d.add_heading("2.1 Kapsam ve Hedefler", level=2)
    d.add_paragraph(
        "Kapsam (içinde): Gezegen seçimi, veri çekme/önbellekleme, stil ve mod "
        "seçenekleri, MIDI/WAV üretimi, canlı önizleme, gelişmiş seçenekler (checkpoint "
        "yolu, LSTM aç/kapa), temel hata yönetimi ve kullanıcı geri bildirimi."
    )
    d.add_paragraph(
        "Kapsam dışı (örnek): Ticari dağıtım, mobil native uygulama, tam profesyonel "
        "DAW entegrasyonu, çok kullanıcılı bulut ölçeği — bu maddeler ileride "
        "genişletme olarak değerlendirilebilir."
    )

    d.add_heading("2.2 Paydaşlar ve Roller", level=2)
    tbl = d.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    rows = [
        ("Paydaş", "Beklenti"),
        ("Danışman öğretim üyesi", "PAT onayı, mimari geri bildirim, vize değerlendirmesi"),
        ("Proje ekibi", "Belgeleme, geliştirme, test, raporlama"),
        ("Son kullanıcı", "Sezgisel arayüz, güvenilir üretim, anlaşılır hata mesajları"),
    ]
    for i, (a, b) in enumerate(rows):
        _set_cell_text(tbl.rows[i].cells[0], a)
        _set_cell_text(tbl.rows[i].cells[1], b)

    d.add_heading("2.3 Çalışma Planı (Örnek Zaman Çizelgesi)", level=2)
    tbl2 = d.add_table(rows=6, cols=3)
    tbl2.style = "Table Grid"
    plan = [
        ("Hafta", "Aşama", "Çıktı"),
        ("1–2", "Planlama", "Kapsam, riskler, araç listesi, görev dağılımı"),
        ("3–4", "Analiz", "Gereksinim listesi, senaryolar, kabul kriterleri"),
        ("5–6", "Tasarım", "Mimari diyagram, API sözleşmesi, veri modeli"),
        ("7–8", "Prototip", "Çalışan uçtan uca akış (MIDI + temel UI)"),
        ("Vize", "Sunum / teslim", "PAT raporu + demo"),
    ]
    for i, row in enumerate(plan):
        for j, val in enumerate(row):
            _set_cell_text(tbl2.rows[i].cells[j], val)

    d.add_heading("2.4 Risk Analizi ve Önlemler", level=2)
    risks = d.add_table(rows=5, cols=3)
    risks.style = "Table Grid"
    rr = [
        ("Risk", "Olasılık / Etki", "Önlem"),
        (
            "Harici API (NASA) kota / kesinti",
            "Orta / Yüksek",
            "Yerel JSON önbelleği, istek gecikmesi (sleep), kullanıcıya bilgi",
        ),
        (
            "ML bağımlılıkları (PyTorch)",
            "Düşük / Orta",
            "Opsiyonel özellik; sunucuda açık hata mesajı, CPU modu",
        ),
        (
            "Ses üretimi (FluidSynth / SoundFont)",
            "Orta",
            "Yapılandırılabilir yol; eksikse sadece MIDI ile devam",
        ),
        (
            "Kapsam genişlemesi",
            "Yüksek",
            "MoSCoW önceliklendirme, vize için PAT odaklı teslim",
        ),
    ]
    for i, row in enumerate(rr):
        for j, val in enumerate(row):
            _set_cell_text(risks.rows[i].cells[j], val)

    d.add_heading("2.5 Kaynaklar ve Araçlar", level=2)
    for item in (
        "Programlama: Python (FastAPI, Uvicorn), TypeScript (React, Vite), üç boyutlu seçici (Three.js / R3F).",
        "Veri: JSON tabanlı gezegen zaman serileri; JSONL eğitim veri setleri.",
        "ML: PyTorch, scikit-learn (isteğe bağlı baseline), checkpoint (.pt, .joblib).",
        "Belgeleme: Markdown / Word; sürüm kontrolü (Git).",
    ):
        d.add_paragraph(item, style="List Bullet")

    d.add_page_break()

    # --- 3. Analiz ---
    d.add_heading("3. Analiz", level=1)

    d.add_heading("3.1 Fonksiyonel Gereksinimler", level=2)
    fr = d.add_table(rows=9, cols=2)
    fr.style = "Table Grid"
    freqs = [
        ("Kod", "Açıklama"),
        ("FR-01", "Kullanıcı gezegen seçebilmeli ve seçim arayüzden API’ye iletilmelidir."),
        ("FR-02", "Sistem, seçilen gezegen için keşif/velocity benzeri veriyi getirmeli veya önbellekten okumalıdır."),
        ("FR-03", "Kullanıcı müzik stili (ör. calm, pop, study, cinematic) seçebilmelidir."),
        ("FR-04", "Üretilen çıktı en az MIDI formatında indirilebilmelidir."),
        ("FR-05", "İsteğe bağlı yüksek kaliteli WAV üretimi yapılandırılabilir olmalıdır."),
        ("FR-06", "Canlı akış / önizleme senaryosu desteklenmelidir (WebSocket veya eşdeğeri)."),
        ("FR-07", "İsteğe bağlı LSTM checkpoint yolu ve etkinleştirme sunucu tarafında çözümlenebilmelidir."),
        ("FR-08", "Hatalar kullanıcıya anlaşılır mesajlarla dönmelidir (ör. eksik SoundFont)."),
    ]
    for i, row in enumerate(freqs):
        for j, val in enumerate(row):
            _set_cell_text(fr.rows[i].cells[j], val)

    d.add_heading("3.2 Fonksiyonel Olmayan Gereksinimler", level=2)
    nfr = d.add_table(rows=6, cols=2)
    nfr.style = "Table Grid"
    nfreq = [
        ("Kod", "Açıklama"),
        ("NFR-01", "Yerel geliştirme ortamında makul sürede yanıt (ör. üretim < birkaç dakika, canlı akış düşük gecikme hedefi)."),
        ("NFR-02", "Yapılandırma .env ile yönetilmelidir (API anahtarı, FluidSynth yolu, checkpoint yolu)."),
        ("NFR-03", "Kod modüler olmalı; servis katmanı (harmoni, sonifikasyon, ML) ayrıştırılmalıdır."),
        ("NFR-04", "Telif ve lisans: Harici MIDI veri kümeleri için lisans bilgisi ayrı dokümanda tutulmalıdır."),
        ("NFR-05", "Güvenlik: API anahtarları repoda düz metin olarak paylaşılmamalıdır."),
    ]
    for i, row in enumerate(nfreq):
        for j, val in enumerate(row):
            _set_cell_text(nfr.rows[i].cells[j], val)

    d.add_heading("3.3 Kullanım Senaryoları (Özet)", level=2)
    for sc in (
        "Senaryo A — Temel üretim: Kullanıcı gezegen ve stil seçer → üret → MIDI indirir.",
        "Senaryo B — Canlı stüdyo: Kullanıcı parametreleri değiştirir → sunucu/istemci akışı ile anlık ses önizlemesi alır.",
        "Senaryo C — Gelişmiş ML: Yönetici .env veya istek parametresi ile LSTM checkpoint tanımlar → melodi katmanı uygulanır.",
        "Senaryo D — Çevrimdışı veri: Önceden kaydedilmiş JSON ile üretim yapılır (API kesintisinde).",
    ):
        d.add_paragraph(sc, style="List Bullet")

    d.add_heading("3.4 Kısıtlar ve Varsayımlar", level=2)
    d.add_paragraph(
        "Kısıtlar: Harici API kullanımında kota; ML eğitimi için donanım süresi; "
        "tek sesli (monofonik) LSTM çıktısının düzenleme motoru ile harmanlanması."
    )
    d.add_paragraph(
        "Varsayımlar: Kullanıcı modern bir tarayıcı kullanır; geliştirici makinesinde "
        "Python 3.10+ ve Node.js kuruludur; ses önizleme için Web Audio API desteklenir."
    )

    d.add_page_break()

    # --- 4. Tasarım ---
    d.add_heading("4. Tasarım", level=1)

    d.add_heading("4.1 Yüksek Seviye Mimari", level=2)
    d.add_paragraph(
        "İstemci (Web SPA) ↔ REST/WebSocket API (FastAPI) ↔ Uygulama servisleri "
        "(veri çekme, sonifikasyon, harmoni, isteğe bağlı LSTM) ↔ Dosya çıktıları "
        "(MIDI/WAV) ve isteğe bağlı ML checkpoint dosyaları."
    )

    d.add_heading("4.2 Modül / Paket Ayrımı (Örnek)", level=2)
    mods = d.add_table(rows=7, cols=2)
    mods.style = "Table Grid"
    modrows = [
        ("Bileşen", "Sorumluluk"),
        ("backend / API", "HTTP uçları, doğrulama, orchestration"),
        ("services / harmony_engine", "Olay (event) üretimi, ölçek, LSTM ile melodi harmanı"),
        ("services / fluid_render", "FluidSynth ile WAV üretimi (yapılandırılabilir)"),
        ("services / lstm_blend", "Checkpoint çözümleme, örnekleme entegrasyonu"),
        ("ml/", "Veri dışa aktarma, eğitim, üretim betikleri"),
        ("web/", "Kullanıcı arayüzü, 3B gezegen seçici, stüdyo sayfaları"),
    ]
    for i, row in enumerate(modrows):
        for j, val in enumerate(row):
            _set_cell_text(mods.rows[i].cells[j], val)

    d.add_heading("4.3 Veri Akışı (Metinsel Diyagram)", level=2)
    d.add_paragraph(
        "Gezegen JSON (points[]) → generate_events (stil, mod, tohum) → olay listesi → "
        "[isteğe bağlı] LSTM ile perde zenginleştirme → MIDI çok kanallı yazım / WAV render."
    )

    d.add_heading("4.4 API ve Veri Modeli (Özet)", level=2)
    d.add_paragraph(
        "İstek/yanıt yapıları JSON olmalıdır. Gezegen noktası örnek alanlar: speed, "
        "zaman, türev/yardımcı özellikler (projeye özgü alan adları dokümantasyonda "
        "sabitlenir). Üretim uçları; stil, mod, tohum, gezegen adı ve isteğe bağlı "
        "checkpoint parametrelerini taşıyabilir."
    )

    d.add_heading("4.5 Güvenlik ve Yapılandırma", level=2)
    d.add_paragraph(
        "Hassas anahtarlar yalnızca sunucu ortamında (.env) tutulur. İstemciye "
        "ham API anahtarı gönderilmez. Üretim ortamında CORS ve rate limiting "
        "değerlendirilir."
    )

    d.add_heading("4.6 Kullanıcı Arayüzü Tasarım Prensipleri", level=2)
    for ui in (
        "Hoş geldin / stüdyo ayrımı; net çağrı-eylem (üret, indir, canlı).",
        "Gezegen seçiminde görsel geri bildirim (3B sahne).",
        "Gelişmiş ayarlar ayrı bölümde; hata ve yükleme durumları kullanıcıya gösterilir.",
    ):
        d.add_paragraph(ui, style="List Bullet")

    d.add_page_break()

    d.add_heading("5. PAT Kapsamında Teslim Özeti", level=1)
    d.add_paragraph(
        "Bu doküman ile Planlama (hedef, zaman planı, risk), Analiz (fonksiyonel ve "
        "fonksiyonel olmayan gereksinimler, senaryolar) ve Tasarım (mimari ve modül "
        "tasarımı) aşamaları vize öncesi tamamlanmış kabul edilebilir. Final teslimde "
        "uygulama (kod), test raporu, kullanım kılavuzu ve sonuç bölümü genişletilecektir."
    )

    d.add_heading("Kaynakça (Örnek)", level=1)
    for ref in (
        "NASA / JPL Horizon Systems (veri kaynağı dokümantasyonu).",
        "FastAPI — https://fastapi.tiangolo.com/",
        "React — https://react.dev/",
        "PyTorch — https://pytorch.org/",
        "MAESTRO veri kümesi lisansı (kullanılıyorsa proje SOURCES_AND_LICENSES dosyasına atıf).",
    ):
        d.add_paragraph(ref, style="List Bullet")

    d.save(out)
    print("Wrote:", out)


if __name__ == "__main__":
    main()
