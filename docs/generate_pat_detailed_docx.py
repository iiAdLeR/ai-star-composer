# -*- coding: utf-8 -*-
"""Detayli PAT (Planlama-Analiz-Tasarim) Word dokumani uretir."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

_BOOKMARK_ID = 1


def set_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str, size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_img_placeholder(doc: Document, title: str, desc: str):
    p = doc.add_paragraph()
    r = p.add_run(f"[BURAYA {title} GORSELINI EKLEYIN]")
    set_font(r, size=11, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(f"Aciklama: {desc}")
    set_font(r, size=10)
    p = doc.add_paragraph()
    r = p.add_run("Sekil basligi: Sekil X. " + title)
    set_font(r, size=10)


def add_figure(doc: Document, image_path: Path, caption: str, width_inch: float = 6.6):
    if image_path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inch))
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return c
    else:
        add_img_placeholder(doc, caption, f"Gorsel bulunamadi: {image_path}")
        return None


def add_bookmark(paragraph, name: str):
    global _BOOKMARK_ID
    bid = _BOOKMARK_ID
    _BOOKMARK_ID += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.append(start)
    paragraph._p.append(end)


def add_internal_link(doc: Document, text: str, anchor: str):
    p = doc.add_paragraph()
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)
    return p


def add_table_2(doc: Document, rows: list[tuple[str, str]], h1="Alan", h2="Detay"):
    t = doc.add_table(rows=len(rows) + 1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = h1
    t.cell(0, 1).text = h2
    for i, (a, b) in enumerate(rows, start=1):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b


def add_table_4(doc: Document, rows: list[tuple[str, str, str, str]], headers=("No", "Is Paketi", "Durum", "Cikti")):
    t = doc.add_table(rows=len(rows) + 1, cols=4)
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.cell(i, j).text = val


def main():
    out = Path(__file__).resolve().parent / "PAT_Detayli_AI_Star_Composer_v8.docx"
    diagrams = Path(__file__).resolve().parent / "diagrams"
    d = Document()

    # Kapak
    add_para(d, "T.C.", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "ISTANBUL TOPKAPI UNIVERSITESI", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "MUHENDISLIK VE DOGA BILIMLERI FAKULTESI", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "YAZILIM MUHENDISLIGI BOLUMU", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()
    add_para(d, "MEZUNIYET PROJESI - PAT RAPORU (DETAYLI)", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "Proje Adi: AI Star Composer", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()
    add_para(d, "Ogrenci 1: AYHAM ELMATAR - 22040301011", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "Ogrenci 2: GAYS HARMUS - 22040301144", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "Danisman: [BURAYA DOKTOR/HOCA ADINI EKLEYIN]", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "Tarih: Nisan 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

    d.add_page_break()

    # Icindekiler (manuel)
    add_para(d, "ICINDEKILER", size=14, bold=True)
    for it in [
        "1. GIRIS VE PROJE OZETI",
        "2. PLANLAMA (PLANLAMA)",
        "3. ANALIZ",
        "4. TASARIM",
        "5. UYGULAMA VE GERCEKLESTIRILEN GELISTIRMELER",
        "6. AKIS DIYAGRAMLARI (DETAYLI)",
        "7. SONUCLAR VE DEGERLENDIRME",
        "8. EKLER (LOG, EKRAN GORUNTULERI, CIKTI TABLOLARI)",
    ]:
        d.add_paragraph(it, style="List Bullet")

    d.add_page_break()

    # 1 GIRIS
    add_para(d, "1. GIRIS VE PROJE OZETI", size=14, bold=True)
    add_para(
        d,
        "AI Star Composer, astronomik verileri muzikal ciktiya donusturen web tabanli bir sistemdir. "
        "Projenin temel hedefi, NASA/JPL gibi kaynaklardan gelen gezegen hareket verilerini "
        "hem bilimsel hem de estetik acidan anlamli bir muzik deneyimine cevirmektir. "
        "Bu kapsamda sistemde veri cekme, veri on isleme, sembolik muzik uretimi, "
        "opsiyonel LSTM tabanli melodi zenginlestirme ve MIDI/WAV cikti alma adimlari uygulanmistir."
    )
    add_para(
        d,
        "Bu rapor, vize sureci icin istenen PAT kapsaminda projenin planlama, analiz ve tasarim "
        "boyutlarini detayli olarak aciklamakta; bunun yaninda gercekte gelistirilen moduller, "
        "egitim ciktilari, test sonuclari ve ekran goruntusu ekleme noktalarini da sistematik sekilde sunmaktadir."
    )
    add_para(
        d,
        "Sistemin veri yolculugu su sekildedir: Kullanici tarafinda secilen gezegen/stil bilgisi "
        "API katmanina ulasir; backend gerekli astronomik noktayi yerel cache veya dis kaynaktan "
        "cozer; harmonic/symbolic motor note event listesi uretir; eger LSTM aktifse model "
        "checkpointinden gelen perde tahminleri event listesine blend edilir; son adimda MIDI/WAV "
        "uretilip istemciye sunulur. Bu akisin her adimi takip edilebilir log ve dosya ciktilariyla "
        "dogrulanmistir."
    )
    cap = add_figure(
        d,
        diagrams / "Ana Ekran.png",
        "Sekil 9. Projenin genel gorunumu (Welcome ana ekrani)",
    )
    if cap is not None:
        add_bookmark(cap, "fig_studio_welcome")

    # 2 PLANLAMA
    d.add_page_break()
    add_para(d, "2. PLANLAMA (PLANLAMA)", size=14, bold=True)
    add_para(d, "2.1 Proje Amaci")
    add_para(
        d,
        "Projenin amaci, kullanicinin secilen bir gezegen baglaminda otomatik muzik uretebilmesini "
        "saglamak ve bu sureci etkileşimli bir web arayuzu uzerinden yonetmektir. "
        "Sistem, veri kaynak bagimliligini azaltmak icin onbellek mekanizmasi kullanir ve "
        "ML entegrasyonu sayesinde melodik zenginlik saglar."
    )
    add_para(d, "2.2 Kapsam")
    add_table_2(
        d,
        [
            ("Kapsam Icinde", "Gezegen secimi, stil secimi, MIDI uretimi, LSTM ile blend, canli/normal stüdyo akislar"),
            ("Kapsam Icinde", "8 gezegen icin veri hazirlama ve JSONL tabanli egitim veri seti olusturma"),
            ("Kapsam Disinda", "Mobil native uygulama gelistirme"),
            ("Kapsam Disinda", "Bulut olcekli dagitim ve ticari lisanslama"),
        ],
    )
    add_para(d, "2.3 Is Paketleri ve Durum")
    add_table_4(
        d,
        [
            ("IP-1", "Gereksinim Analizi ve Planlama", "Tamamlandi", "Gereksinim listesi, PAT taslagi"),
            ("IP-2", "Astronomik Veri Toplama", "Tamamlandi", "data/nasa_planets/*.json"),
            ("IP-3", "Veri On Isleme ve Donusum", "Tamamlandi", "style_sequences_8planets.jsonl"),
            ("IP-4", "Sembolik Muzik Uretim Altyapisi", "Tamamlandi", "harmony_engine entegrasyonu"),
            ("IP-5", "LSTM Egitimi ve Entegrasyonu", "Tamamlandi", "note_lstm_style_planet.pt"),
            ("IP-6", "Web Arayuzu ve Etkilesim", "Tamamlandi", "PlanetPicker3D + Studio/Live sayfalari"),
            ("IP-7", "Test ve Cikti Dogrulama", "Tamamlandi", "final MIDI seti + loglar"),
        ],
    )
    add_para(d, "2.4 Riskler ve Onlemler")
    add_table_2(
        d,
        [
            ("API gecikmesi / kota", "Yerel cache kullanimi (data/nasa_planets)"),
            ("Model bagimlilik sorunlari", "ML modullerini opsiyonel tasarlama, checkpoint kontrolu"),
            ("Render bagimliligi", "SoundFont yoksa sadece MIDI cikti ile devam"),
            ("Kapsam buyumesi", "PAT odakli onceliklendirme ve parcali teslim"),
        ],
    )
    add_para(
        d,
        "Planlama asamasinda her is paketi bagimlilik sirasina gore tanimlanmistir. "
        "Ornegin IP-3 (veri donusum) IP-2 tamamlanmadan baslamaz; IP-5 (LSTM) ise "
        "IP-3 ciktilari olmadan ilerleyemez. Bu bagimliliklar sure kaymalarini azaltmis "
        "ve vize odakli teslimi guvenceye almistir."
    )
    add_figure(
        d,
        diagrams / "gantt_is_paketi.png",
        "Sekil 7. Is paketi zaman cizelgesi (Gantt)",
    )

    # 3 ANALIZ
    d.add_page_break()
    add_para(d, "3. ANALIZ", size=14, bold=True)
    add_para(d, "3.1 Fonksiyonel Gereksinimler")
    add_table_2(
        d,
        [
            ("FR-01", "Kullanici gezegen secimi yapabilmelidir."),
            ("FR-02", "Sistem secilen gezegen verisini cekmeli veya cache'ten okumalidir."),
            ("FR-03", "Kullanici stil secimi yapabilmelidir (calm/pop/study/cinematic)."),
            ("FR-04", "Sistem MIDI cikti olusturabilmelidir."),
            ("FR-05", "LSTM checkpoint aktif oldugunda melodi blend uygulanabilmelidir."),
            ("FR-06", "Canli akis senaryosu desteklenmelidir."),
            ("FR-07", "Ciktilar kullaniciya indirilebilir formatta sunulmalidir."),
        ],
    )
    add_para(d, "3.2 Fonksiyonel Olmayan Gereksinimler")
    add_table_2(
        d,
        [
            ("NFR-01", "Kullanilabilirlik: Arayuz anlasilir ve akis net olmalidir."),
            ("NFR-02", "Performans: Muzik uretimi makul surede tamamlanmalidir."),
            ("NFR-03", "Modulerlik: Servis katmanlari ayrik ve bakimi kolay olmalidir."),
            ("NFR-04", "Yapilandirma: .env ile degisken yonetimi uygulanmalidir."),
            ("NFR-05", "Guvenlik: Hassas anahtarlar kod disinda saklanmalidir."),
        ],
    )
    add_para(d, "3.3 Kullanici Senaryolari")
    for txt in [
        "Senaryo-1 (Temel): Kullanici gezegen + stil secer, sistem MIDI uretir.",
        "Senaryo-2 (Gelismis): Kullanici LSTM aktif eder, model tabanli melodi zenginlestirme uygulanir.",
        "Senaryo-3 (Canli): Kullanici Studio Live ekraninda anlik akisi izler/yonlendirir.",
        "Senaryo-4 (Offline): Sistem internet olmasa da cache JSON ile uretim yapar.",
    ]:
        d.add_paragraph(txt, style="List Bullet")
    add_figure(
        d,
        diagrams / "use_case_diagram.png",
        "Sekil 6. Kullanici-sistem etkilesimini gosteren Use Case diyagrami",
    )
    add_para(
        d,
        "Analiz sonucunda kritik kabul kriterleri su sekilde netlestirilmistir: "
        "(1) Gezegen secimi muzikal ciktiyi degistirmelidir, (2) Stil secimi tonal/ritmik "
        "karakteri degistirmelidir, (3) LSTM acik/kapali durumunda duyulabilir fark elde "
        "edilmelidir, (4) Uretilen dosya indirilebilir ve tekrar oynatilabilir olmalidir."
    )

    # 4 TASARIM
    d.add_page_break()
    add_para(d, "4. TASARIM", size=14, bold=True)
    add_para(d, "4.1 Mimari Tasarim (Yuksek Seviye)")
    add_para(
        d,
        "Mimari; istemci (React tabanli web), API katmani (FastAPI), servis katmani "
        "(harmony_engine, sonifier, lstm_blend), veri katmani (JSON/JSONL/checkpoint), "
        "ve cikti katmani (MIDI/WAV) olarak bolunmustur."
    )
    add_figure(
        d,
        diagrams / "architecture_overview.png",
        "Sekil 1. AI Star Composer yuksek seviye mimari diyagrami",
    )
    add_para(d, "4.2 Modul Tasarimi")
    add_table_2(
        d,
        [
            ("web/src/pages", "Studio, StudioLive, Welcome ekranlari"),
            ("web/src/components/planets", "PlanetPicker3D ve ilgili gorsel bilesenler"),
            ("backend/api.py", "REST/WS endpoint yonetimi"),
            ("services/harmony_engine.py", "Sembolik event uretimi ve blend islemleri"),
            ("services/lstm_blend.py", "Checkpoint cozumu, planet/style indexleme"),
            ("ml/train_sequence_lstm.py", "LSTM egitimi (style+planet destekli)"),
            ("ml/generate_from_lstm.py", "Checkpoint'ten MIDI uretimi"),
        ],
    )
    add_figure(
        d,
        diagrams / "module_relationship_diagram.png",
        "Sekil 8. Sinif/modul iliski diyagrami (bilesen baglantilari)",
    )
    add_para(d, "4.3 Veri Tasarimi")
    add_table_2(
        d,
        [
            ("Ham veri", "data/nasa_planets/*.json"),
            ("Egitim veri seti", "data/ml/style_sequences_8planets.jsonl"),
            ("LSTM checkpoint", "ml/checkpoints/note_lstm_style_planet.pt"),
            ("Baseline model", "ml/checkpoints/pitch_rf.joblib"),
            ("Final cikti", "outputs/*_final_*_lstm_symphony_planet.mid"),
        ],
    )
    add_para(
        d,
        "Veri tasariminda dosya bazli bir model benimsenmistir: ham astronomik JSON dosyalari "
        "egitim asamasinda etiketli JSONL formatina donusur; egitim sonunda checkpoint dosyalari "
        "uretilir; uretim asamasinda bu checkpointler API/ML scriptleri tarafindan okunup "
        "MIDI ciktisi olusturur. Bu yapi, veritabanina bagimli olmadan tekrar uretilebilir "
        "pipeline saglamaktadir."
    )
    add_figure(
        d,
        diagrams / "ml_pipeline_flow.png",
        "Sekil 2. Veri donusum ve ML pipeline akis diyagrami",
    )

    # 5 UYGULAMA
    d.add_page_break()
    add_para(d, "5. UYGULAMA VE GERCEKLESTIRILEN GELISTIRMELER", size=14, bold=True)
    add_para(d, "5.1 Veri Hazirlama ve Birlestirme")
    add_para(
        d,
        "8 gezegen verisi data/nasa_planets klasorunde toplanmis ve ml.bundle_local_planet_exports "
        "betigi ile tek bir JSONL dosyasina birlestirilmistir. Varsayilan tohum araligi 0..31 olarak "
        "duzenlenmis, toplam 2048 satirlik bir egitim veri seti elde edilmistir."
    )
    add_para(d, "5.2 LSTM Egitimi (Style + Planet)")
    add_para(
        d,
        "ml.train_sequence_lstm betigi style_idx ve planet_idx alanlarini otomatik algilayacak "
        "sekilde kullanilmistir. 45 epoch sonunda note_lstm_style_planet.pt dosyasi uretilmistir. "
        "Egitim loglarinda kayip degerinin belirgin sekilde dustugu gozlenmistir."
    )
    add_para(d, "5.3 Baseline Model")
    add_para(
        d,
        "train_baseline_sklearn ile RandomForest tabanli baseline model egitilmis "
        "ve pitch_rf.joblib dosyasi olusturulmustur. Bu model karsilastirma amacli saklanmistir."
    )
    add_para(d, "5.4 Harici Veri (MAESTRO) Denemesi")
    add_para(
        d,
        "external_notes_maestro.jsonl uzerinden style/planet kosulu kapali olacak sekilde "
        "ayri bir LSTM egitimi yapilarak note_lstm_external_maestro.pt elde edilmistir."
    )
    cap = add_figure(
        d,
        diagrams / "Egitim Log Ekran Goruntusu.png",
        "Sekil 10. Terminal egitim log ekran goruntusu",
    )
    if cap is not None:
        add_bookmark(cap, "fig_training_log")
    cap = add_figure(
        d,
        diagrams / "Checkpoint Dosyalari Ekran Goruntusu.png",
        "Sekil 11. Checkpoint dosyalari ekran goruntusu",
    )
    if cap is not None:
        add_bookmark(cap, "fig_checkpoints")
    add_para(
        d,
        "Uygulama asamasinda dosya isimlendirme standardi da tanimlanmistir: "
        "planet_final_planet_style_lstm_symphony_planet.mid. Bu standard sayesinde "
        "raporlamada style/planet eslestirmesi otomatik cikartilabilmistir."
    )

    add_para(d, "5.5 Astronomik Veri Nasil Toplandi? (Detayli)", bold=True)
    add_para(
        d,
        "Veri toplama katmani scripts/data_fetcher.py dosyasinda uygulanmistir. "
        "PLANET_IDS sabiti ile Mercury(199) - Neptune(899) hedef kodlari sabitlenmistir. "
        "Sistem, Horizons API'ye EPHEM_TYPE=VECTORS parametresiyle gunluk adimli istek atar."
    )
    for txt in [
        "Adim-1: Kullanici gezegen adini secer (ornegin Mars).",
        "Adim-2: _resolve_target fonksiyonu ad -> NASA target kodu eslesmesi yapar.",
        "Adim-3: _build_date_range ile baslangic ve bitis tarihi hesaplanir.",
        "Adim-4: _request_horizons ile JSON yaniti cekilir.",
        "Adim-5: _extract_vectors regex ile X,Y,Z,VX,VY,VZ ve tarih satirlarini ayiklar.",
        "Adim-6: Her satirdan speed, radius, radial_velocity, light_intensity_proxy, heading_xy, speed_delta turetilir.",
        "Adim-7: Dataset metadata + points[] yapisinda kaydedilir (save_velocity_dataset).",
    ]:
        d.add_paragraph(txt, style="List Bullet")
    add_para(
        d,
        "Bu mekanizma sayesinde veri nereden geliyor sorusunun cevabi nettir: "
        "NASA Horizons -> requests -> parse -> normalized point list -> JSON cache."
    )

    add_para(d, "5.6 3B Gezegenler Nasil Yapildi? (PlanetPicker3D Detayi)", bold=True)
    add_para(
        d,
        "3B katman web/src/components/planets/PlanetPicker3D.tsx dosyasinda React Three Fiber "
        "ile yazilmistir. Canvas icerisinde Scene bileseni ortam isiklari, yildiz alani ve "
        "kamera riglerini yonetir."
    )
    for txt in [
        "Gezegen geometrisi: PlanetSurface icerisinde sphereGeometry + meshStandardMaterial kullanilir.",
        "Doku (texture): PLANET_TEXTURE_URL ile URL secilir, useTexture ile yuklenir, SRGB colorSpace uygulanir.",
        "Gezegen konumlari: PlanetCarousel icinde acisal dagitim ile halka (ring) uzerine yerlestirilir.",
        "Donme (rotation): PlanetBody useFrame dongusunde rotation.y her frame artirilir.",
        "Secili gezegen efekti: scale SELECTED_SCALE_MUL ile buyutulur, bob animasyonu uygulanir.",
        "Saturn halkasi: SaturnRings bileseninde torusGeometry ile render edilir.",
        "Dunya-ayi gorevi: EarthMoonMission bileseni CatmullRomCurve3 yolu uzerinde roket animasyonu cizer.",
    ]:
        d.add_paragraph(txt, style="List Bullet")

    add_para(d, "5.7 Gezegeni Dondurme ve Kamera Modlari", bold=True)
    add_para(
        d,
        "Kullanici deneyiminde uc kamera modu vardir: static, showcase ve orbit. "
        "Showcase modunda kamera secili gezegen etrafinda sinematik sekilde yumusak hareket eder. "
        "Orbit modunda OrbitControls aktif olur, kullanici fare ile etrafinda doner."
    )
    for txt in [
        "Drag mekanigi: onPointerDown -> pointermove ile dragTwist guncellenir.",
        "Snap mekanigi: snapFromDrag fonksiyonu aciyi en yakin gezegene yuvarlar.",
        "Klavye cikisi: orbit modunda ESC ile showcase moduna donulur.",
        "Kalite optimizasyonu: PerformanceMonitor fps'e gore low/medium/high kalite secimini otomatik ayarlar.",
    ]:
        d.add_paragraph(txt, style="List Bullet")

    add_para(d, "5.8 Sembolik Muzik Uretimi Nasil Calisiyor?", bold=True)
    add_para(
        d,
        "Temel muzikal omurga services/harmony_engine.py dosyasindaki generate_events fonksiyonudur. "
        "Bu fonksiyon her point kaydini muzik event'ine cevirir. Event yapisinda time, duration, "
        "base_note, lead_note, bass_note, harmony, velocity ve pan bulunur."
    )
    for txt in [
        "Speed normalization -> scale index map ile temel nota secimi.",
        "Style etkisi: calm/pop/study/cinematic icin farkli ritim, leap, velocity kurallari.",
        "Planet etkisi: get_planet_style_rhythm ve get_planet_style_voice ile style davranisi gezegene gore tilt edilir.",
        "Anti-repetition: repeat memory, melodic leap constraint, metric accent uygulamalari.",
        "Quantize: zaman ve sureler grid'e snap edilerek muziksel duzen korunur.",
    ]:
        d.add_paragraph(txt, style="List Bullet")

    add_para(d, "5.9 LSTM Entegrasyonu ve Egitim Detayi", bold=True)
    add_para(
        d,
        "LSTM boru hatti ml/export_style_sequences.py -> ml/bundle_local_planet_exports.py -> "
        "ml/train_sequence_lstm.py -> ml/generate_from_lstm.py sirasiyla ilerler."
    )
    for txt in [
        "export_style_sequences: her satirda notes[] + style_idx + planet_idx etiketleri yazilir.",
        "bundle_local_planet_exports: data/nasa_planets altindaki tum JSON dosyalarini tek JSONL'de birlestirir.",
        "train_sequence_lstm: style_idx/planet_idx varligini auto tespit eder; NoteLSTM embedding katmanlarini ona gore acar.",
        "Checkpoint icerigi: model state + seq_len + use_style + use_planet + num_styles + num_planets.",
        "generate_from_lstm: style_name_to_idx ve planet_name_to_idx ile kosullu sampling yapar.",
        "services/lstm_blend: API/canli akis sirasinda checkpoint varsa olay notasini LSTM ile blend eder.",
    ]:
        d.add_paragraph(txt, style="List Bullet")

    add_para(d, "5.10 API'de Veri Nereden Nereye Gidiyor? (Cagri Zinciri)", bold=True)
    add_para(
        d,
        "Servis zinciri services/generation_service.py ve services/live_stream_service.py dosyalarinda "
        "somutlasmistir. Uretim isteginde akis su sekildedir:"
    )
    for txt in [
        "API istegi -> fetch_velocity_dataset (NASA veya cache veri noktasi).",
        "points -> generate_note_events/generate_events (sembolik event listesi).",
        "Opsiyonel: apply_lstm_checkpoint_to_events (checkpoint aktifse pitch blend).",
        "save_symphony_midi_from_events ile cok kanalli MIDI yazimi.",
        "render_events_to_wav ve opsiyonel FluidSynth ile HQ WAV.",
        "Sonuc JSON'unda data_json, midi, wav, metrics ve lstm_meta donusu.",
    ]:
        d.add_paragraph(txt, style="List Bullet")

    add_para(d, "5.11 Dosya Bazli Ispat (Traceability)", bold=True)
    add_table_4(
        d,
        [
            ("T-1", "Ham veri kaynagi", "scripts/data_fetcher.py", "NASA Horizons cevabinin parse edilmesi"),
            ("T-2", "3B arayuz", "PlanetPicker3D.tsx", "Gezegen render, donus, kamera ve drag"),
            ("T-3", "Sembolik motor", "services/harmony_engine.py", "Event bazli nota/harmony uretimi"),
            ("T-4", "LSTM egitim", "ml/train_sequence_lstm.py", "Style+planet kosullu model egitimi"),
            ("T-5", "LSTM uretim", "ml/generate_from_lstm.py", "Checkpointten note sampling"),
            ("T-6", "API entegrasyon", "services/generation_service.py", "Uctan uca cikti dosyasi olusturma"),
        ],
        headers=("No", "Bilesen", "Dosya", "Aciklama"),
    )

    # 6 AKIS DIYAGRAMLARI
    d.add_page_break()
    add_para(d, "6. AKIS DIYAGRAMLARI (DETAYLI)", size=14, bold=True)
    add_para(d, "6.1 Uctan Uca Sistem Akisi")
    add_para(
        d,
        "Kullanici arayuzde gezegen ve stil secer -> istek API'ye gider -> veri kaynagi "
        "cache/API olarak cozulur -> sembolik event uretimi yapilir -> (opsiyonel) LSTM blend "
        "uygulanir -> MIDI yazilir -> kullaniciya sonuc donulur."
    )
    add_figure(
        d,
        diagrams / "end_to_end_flow.png",
        "Sekil 3. End-to-end sistem akis diyagrami",
    )
    add_para(d, "6.2 ML Egitim Akisi")
    add_para(
        d,
        "Planet JSON dosyalari -> export/bundle -> style_sequences_8planets.jsonl -> "
        "train_sequence_lstm -> checkpoint (.pt) -> generate_from_lstm/symphony_from_lstm."
    )
    add_figure(
        d,
        diagrams / "ml_pipeline_flow.png",
        "Sekil 4. ML egitim akisi (planet JSON -> JSONL -> checkpoint -> MIDI)",
    )
    add_para(d, "6.3 Canli (Live) Akis")
    add_para(
        d,
        "Studio Live ekraninda kullanici parametre degisikligi yapar -> WebSocket/istek "
        "sunucuya iletilir -> guncel event/melodi hesaplanir -> istemciye geri donus "
        "saglanir -> kullanici anlik geri bildirim alir."
    )
    add_figure(
        d,
        diagrams / "live_sequence.png",
        "Sekil 5. Studio Live sequence diyagrami",
    )
    add_para(
        d,
        "Bu diyagramda her mesaj adimi tek tek tanimlanmistir: istemci secim yapar, "
        "Web UI payload olusturur, API dogrular, servis katmani event+blend hesaplar, "
        "sonuc tekrar istemciye akar. Boylece 'veri nereden geliyor ve nereye gidiyor' "
        "sorusu operasyonel olarak net cevaplanmis olur."
    )

    # 7 SONUC
    d.add_page_break()
    add_para(d, "7. SONUCLAR VE DEGERLENDIRME", size=14, bold=True)
    add_para(d, "7.1 Somut Ciktilar")
    for txt in [
        "8 gezegenli egitim veri seti (2048 satir) olusturuldu.",
        "Style + Planet kosullu LSTM modeli egitildi (note_lstm_style_planet.pt).",
        "Baseline RandomForest modeli egitildi (pitch_rf.joblib).",
        "MAESTRO tabanli ek LSTM modeli egitildi (note_lstm_external_maestro.pt).",
        "Mars/Jupiter icin 4 stilde toplam 8 adet final symphony MIDI dosyasi uretildi.",
    ]:
        d.add_paragraph(txt, style="List Bullet")
    add_para(d, "7.2 Teknik Degerlendirme")
    add_para(
        d,
        "Elde edilen sonuclar, sistemin planlanan PAT hedeflerini asarak calisir bir "
        "prototip seviyesine ulastigini gostermektedir. Moduler yapi sayesinde hem "
        "kural tabanli hem de ML destekli muzik uretimi birlikte kullanilmistir. "
        "Arayuz tarafinda 3B gezegen secici ve canli akis deneyimi, kullanici etkilesimini "
        "guclendirmistir."
    )
    add_para(d, "7.3 Gelecek Calismalar")
    for txt in [
        "Transformer tabanli model ile daha uzun yapisal tutarlilik denemeleri",
        "Daha fazla gezegen/astronomik parametre ile dataset genisletme",
        "Kullaniciya ozel stil/duygu kosullandirmasi",
        "Bulut tabanli dagitim ve daha kapsamli performans testleri",
    ]:
        d.add_paragraph(txt, style="List Bullet")
    cap = add_figure(
        d,
        diagrams / "stdio.png",
        "Sekil 12. Studio sayfasi goruntusu",
    )
    if cap is not None:
        add_bookmark(cap, "fig_studio_page")
    add_table_4(
        d,
        [
            ("R-1", "JSONL satir sayisi", "2048", "8 gezegen x 4 stil x 2 mod x 32 seed kombinasyonu"),
            ("R-2", "LSTM ana checkpoint", "note_lstm_style_planet.pt", "Style + planet kosullu egitim"),
            ("R-3", "Baseline model", "pitch_rf.joblib", "Karsilastirma amacli saklandi"),
            ("R-4", "Final MIDI sayisi", "8 adet", "Mars ve Jupiter icin 4 stil seti"),
        ],
        headers=("MetriK", "Parametre", "Deger", "Yorum"),
    )

    # 8 EKLER
    d.add_page_break()
    ekler_head = add_para(d, "8. EKLER", size=14, bold=True)
    add_bookmark(ekler_head, "ekler_start")
    add_para(d, "Ek-1: Egitim Loglari", bold=True)
    add_para(d, "Model egitimi sirasinda kullanilan komutlar ve terminal ciktilari bu ekte sunulmustur.")
    add_internal_link(d, "Ilgili sekle git: Sekil 10 (Terminal egitim logu)", "fig_training_log")
    add_para(d, "Ek-2: Output Dosya Tablosu", bold=True)
    add_para(d, "Final uretilen MIDI dosyalarinin planet/style eslestirme tablosu bu ekte verilmisitir.")
    add_internal_link(d, "Ilgili sekle git: Sekil 11 (Checkpoint goruntusu)", "fig_checkpoints")
    add_para(d, "Ek-3: Arayuz Ekran Goruntuleri", bold=True)
    add_para(d, "Arayuz ekran goruntuleri ve cekim kontrol listesi bu ekte bulunmaktadir.")
    add_internal_link(d, "Ilgili sekle git: Sekil 9 (Welcome goruntusu)", "fig_studio_welcome")
    add_internal_link(d, "Ilgili sekle git: Sekil 12 (Studio goruntusu)", "fig_studio_page")

    # Son kontrol notu
    d.add_paragraph()
    add_para(
        d,
        "NOT: Final teslim oncesi sekil numaralari (Sekil 1..12) ve danisman bilgisi kontrol edilmelidir.",
        size=10,
        bold=True,
    )

    d.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()

